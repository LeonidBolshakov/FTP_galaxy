from sys import argv

from pathlib import Path
from docx import Document
from docx.oxml.ns import nsdecls
from docx.enum.dml import MSO_THEME_COLOR
import docx

import constant as C


class MyException(Exception):
    def __init__(self, text_err, ret_code):
        self.text_err = text_err
        self.ret_code = ret_code


def main():
    """ Функция main():
            Формирует WORD файл для рассылки сообщений о выходе обновлений компонент.
    return:
            None
    """
    check_options()

    try:
        release_date = input('Введите дату выпуска обновлений:\n').strip()
    except KeyboardInterrupt:
        raise MyException('Оператор прекратил работу программы', 777)

    doc = Document()
    paragraph = doc.add_paragraph()

    try:
        with open(C.TEXT_MAIL) as file_body_mail:
            body_mail = file_body_mail.readlines()
    except FileNotFoundError:
        raise MyException(f'не могу открыть шаблон письма "{C.TEXT_MAIL}" ', 1000)

    for line_mail in body_mail:
        line_mail = line_mail.split('=', 2)
        lm_key = line_mail[0].strip().lower()
        lm_value = line_mail[1].lstrip()
        match lm_key:
            case 'текст': add_text(paragraph, lm_value)
            case 'дата': add_text(paragraph, release_date)
            case 'текстд': add_text(paragraph, processing_ТекстД_template(lm_value, release_date))
            case 'ссылка': add_hyperlink(paragraph, lm_value, lm_value)
            case 'компоненты': add_text(paragraph, get_list_components())
            case _: raise MyException(f'В шаблоне письма нераспознанный ключ строки - "{line_mail[0]}"', 777)

    try:
        doc.save(C.WORD_NAME)
    except PermissionError as err:
        raise MyException(f'Нет доступа к файлу "{C.WORD_NAME}"\n'
                          f'{err}', 777)

    input('Нажмите Enter')


def check_options():
    """ Функция check_options():
            Проверяет количество переданных программе параметров.
        return:
            None
    """
    if len(argv) != 2:
        raise MyException("Программе передано неверное количество параметров.\n"
                          "Программа принимает 1 параметр:\n"
                          "   Имя директории с локальным дистрибутивом обновлений, например,\n"
                          "   c:\\Дистрибутив\\PREPARE').", 1000)


def add_text(paragraph: docx.text.paragraph, txt: str) -> None:
    """ Функция add_text(paragraph: docx.text.paragraph, txt: str) -> None:
            добавляет текст в WORD файл
        :param
            paragraph:   Параграф (абзац), в который добавляется текст.
            txt:         Добавляемый текст
        :return:
            None
    """
    txt = txt.replace(r'\n', '\n')
    run = paragraph.add_run(txt)
    run.font.name = C.WORD_FONT_NAME
    run.font.size = C.WORD_FONT_SIZE


def processing_ТекстД_template(lm_value: str, release_date: str) -> str:
    """ Функция processing_ТекстД_template(lm_value: str, release_date: str) -> str:
            Функция обрабатывает строку шаблона 'ТекстД'
        :param
            1. lm_value:        - обрабатываемая строка
            2. release_date:    - дата выпуска компонент
        :return
            Обработанная строка
    """
    lm_value = lm_value.replace('\n', ' ')
    return ' ' + lm_value if release_date else lm_value[0].upper() + lm_value[1:]


def add_hyperlink(paragraph: docx.text.paragraph, txt: str, url: str) -> None:
    """ Функция add_hyperlink(paragraph: docx.text.paragraph, txt: str, url: str) -> None:
            Добавляет внешнюю гиперссылку в WORD файл
        :param
            1. paragraph:   Параграф (абзац), в который добавляется гиперссылка
            2. txt:         Текстовое отображение гиперссылки.
            3. url:         Внешняя гиперссылка
        :return:
            None
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = txt
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.theme_color = MSO_THEME_COLOR.HYPERLINK
    r.font.underline = True
    r.font.name = C.WORD_FONT_NAME
    r.font.size = C.WORD_FONT_SIZE


def get_list_components() -> str:
    """ Функция get_list_components() -> str:
            Из списка файлов поддиректории NEW формирует состав выпущенных компонет.
        return:
            Строка, содержащая список файлов, выпущенных компонент.
    """
    dir_components = Path(argv[1])
    sub_dir_new = Path(dir_components, C.SUB_DIR_NEW)

    if sub_dir_new.is_dir():
        component_files = sub_dir_new.glob('*.*')
    else:
        raise MyException(f'Нет доступа к каталогу {sub_dir_new}\n'
                          f'Обратитесь к системному администратору.', 1000)

    result = ''
    for file in component_files:
        result += name_file_to_component(file) + C.COMPONENT_SEPARATOR

    return processing_end_of_line(result)


def name_file_to_component(file: Path) -> str:
    """ Функция name_file_to_component(file: Path) -> str:
            Преобразует имя файла в имя компонента.
        :param
            file: Имя файла (Name.Ext)
        :return:
            Имя компонента (Name)
    """
    return file.stem


def processing_end_of_line(line: str) -> str:
    """ Функция processing_end_of_line(line: str) -> str:
            Убирает из строки разделитель после последнего элемента и формирует новый  конец строки
        param
            line: строка со списком компонент.
        return:
            Обработанная строка со списком компонент.
    """
    return line[:-len(C.COMPONENT_SEPARATOR)] + '.\n'


def processing_error() -> None:
    """ Функция def processing_error() ->:
            Завершает работу программы после генерации сообщения об ошибке.
        return:
            None
    """
    print(f'{e.text_err}\n')
    print(f'Программа закончила свою работу\n'
          f'Код возврата {e.ret_code}')
    exit(e.ret_code)


if __name__ == '__main__':

    try:
        main()
    except MyException as e:
        processing_error()
    else:
        print('Программа закончила свою работу\n'
              'Код возврата - 0')

