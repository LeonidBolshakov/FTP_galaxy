import sys
from os import startfile

from docx import Document
from docx.shared import Cm, Inches
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt

from constant import const as c, ColumnProperties


class OneSectionWordDocument:
    """Класс для работы с документом Word (одна секция)."""

    def __init__(self):
        self.doc = Document()  # Объект документа
        self.section = self.doc.sections[0]  # Первая (единственная) секция

    def save(self, name_document: str) -> None:
        """Сохраняет документ. Обрабатывает ошибку доступа."""
        try:
            self.doc.save(name_document)
        except PermissionError:
            sys.tracebacklimit = 0
            raise PermissionError(c.TEXT_ERROR_SAVE) from None

    def set_orientation(self, orientation: WD_ORIENTATION) -> None:
        """Устанавливает ориентацию страницы."""
        self.section.orientation = orientation
        # Меняет ширину и высоту местами
        self.section.page_height, self.section.page_width = (
            self.section.page_width,
            self.section.page_height,
        )

    def set_margins(
        self, left: c.Size, right: c.Size, top: c.Size, bottom: c.Size
    ) -> None:
        """Устанавливает поля страницы."""
        self.section.left_margin = left
        self.section.right_margin = right
        self.section.bottom_margin = bottom
        self.section.top_margin = top


class TableDocWord:
    """Класс для работы с таблицами в Word."""

    def __init__(self, doc, properties_columns: list[ColumnProperties]) -> None:
        self.doc = doc
        self.properties_columns = properties_columns  # Свойства столбцов
        self.table = doc.add_table(rows=1, cols=len(properties_columns))
        self.set_head()  # Создание заголовка

    def set_width_column(self, num_column: int, width: c.Size) -> None:
        """Устанавливает ширину столбца."""
        self.table.allow_autofit = False
        self.table.columns[num_column].width = width
        # Применяет ширину ко всем ячейкам столбца
        for cell_column in self.table.columns[num_column].cells:
            cell_column.width = width

    def set_head(self) -> None:
        """Добавляет заголовки таблицы."""
        head_cells = self.table.rows[0].cells
        for i, item in enumerate(self.properties_columns):
            p = head_cells[i].paragraphs[0]
            p.add_run(item.head_text)  # Текст заголовка
            self.set_width_column(i, item.head_width)  # Ширина столбца

    def set_table_style(self, style: str) -> None:
        """Устанавливает стиль таблицы"""
        self.table.style = style

    def add_line(self, items: list):
        """Добавляет строку с данными в таблицу."""
        cells = self.table.add_row().cells
        for i, item in enumerate(items):
            cells[i].text = str(item)
            # Настройка шрифта для ячейки
            paragraph = cells[i].paragraphs[0]
            paragraph.runs[0].font.name = self.properties_columns[i].font_name
            paragraph.runs[0].font.size = self.properties_columns[i].font_size
            paragraph.runs[0].font.bold = self.properties_columns[i].font_bold


def test():
    items_row = (
        [7, "1024", "Плюшевые котята", "1", "2"],
        [3, "2042", "Меховые пчелы", "3", "4"],
        [1, "1288", "Ошейники для пуделей", "5", "6"],
    )

    properties_columns = [
        ColumnProperties("Первый", Cm(1), None, None, font_bold=False),
        ColumnProperties("Второй", Cm(2), "Arial", Pt(15), font_bold=True),
        ColumnProperties("Третий", Inches(3), "Calibri", None),
        ColumnProperties("Четвёртый", Cm(4), "Ink Free", Pt(12)),
        ColumnProperties("Пятый", Cm(5)),
    ]

    doc_main = OneSectionWordDocument()
    doc_main.set_orientation(WD_ORIENTATION.LANDSCAPE)
    doc_main.set_margins(Cm(1), Cm(1), Cm(1), Cm(1))
    table = TableDocWord(doc=doc_main.doc, properties_columns=properties_columns)
    table.set_table_style("Light Grid Accent 1")
    for item_row in items_row:
        table.add_line(item_row)
    file_docx = "test.docx"
    doc_main.save(file_docx)
    startfile(file_docx)


if __name__ == "__main__":
    test()
